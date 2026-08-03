#!/usr/bin/env python3
"""Extract one DOCX table into a UTF-8 CSV file."""

from __future__ import annotations

import argparse
import csv
from pathlib import Path

from docx import Document


def main() -> int:
    """Run the table extractor.

    Returns:
        Process exit code.
    """
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="Source DOCX")
    parser.add_argument("output", type=Path, help="Destination CSV")
    parser.add_argument("--table", type=int, default=0, help="Zero-based table index")
    args = parser.parse_args()
    input_path = args.input.expanduser().resolve(strict=False)
    output_path = args.output.expanduser().resolve(strict=False)
    if not input_path.is_file():
        parser.error(f"Input file not found: {input_path}")
    if output_path.suffix.lower() != ".csv":
        parser.error("Output path must end with .csv")
    if args.table < 0:
        parser.error("Table index must be non-negative.")
    try:
        document = Document(input_path)
        table = document.tables[args.table]
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with output_path.open("w", encoding="utf-8-sig", newline="") as output_file:
            writer = csv.writer(output_file)
            for row in table.rows:
                writer.writerow([cell.text for cell in row.cells])
    except IndexError:
        parser.error(
            f"Table index {args.table} is out of range; document has {len(document.tables)} table(s)."
        )
    except Exception as exc:
        parser.error(str(exc))
    print(f"created: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
