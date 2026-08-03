#!/usr/bin/env python3
"""Inspect DOCX structure, content features, and metadata as JSON."""

from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_REL = "{http://schemas.openxmlformats.org/package/2006/relationships}"


def inspect_docx(path: Path, sample_paragraphs: int) -> dict:
    """Inspect a DOCX package.

    Args:
        path: DOCX input.
        sample_paragraphs: Maximum non-empty paragraph samples.

    Returns:
        JSON-serializable inspection data.
    """
    document = Document(path)
    paragraphs = [
        paragraph for paragraph in document.paragraphs if paragraph.text.strip()
    ]
    headings = [
        {"style": paragraph.style.name, "text": paragraph.text}
        for paragraph in paragraphs
        if paragraph.style and paragraph.style.name.startswith("Heading")
    ]
    tables = [
        {
            "index": index,
            "rows": len(table.rows),
            "columns": max((len(row.cells) for row in table.rows), default=0),
            "sample": [[cell.text for cell in row.cells[:8]] for row in table.rows[:4]],
        }
        for index, table in enumerate(document.tables)
    ]
    sections = [
        {
            "index": index,
            "page_width_emu": section.page_width,
            "page_height_emu": section.page_height,
            "orientation": str(section.orientation),
            "top_margin_emu": section.top_margin,
            "right_margin_emu": section.right_margin,
            "bottom_margin_emu": section.bottom_margin,
            "left_margin_emu": section.left_margin,
        }
        for index, section in enumerate(document.sections)
    ]
    properties = document.core_properties

    features = {
        "tracked_insertions": 0,
        "tracked_deletions": 0,
        "content_controls": 0,
        "simple_fields": 0,
        "field_instructions": 0,
        "comment_references": 0,
    }
    external_links: list[str] = []
    with zipfile.ZipFile(path) as package:
        names = package.namelist()
        for name in names:
            if name.startswith("word/") and name.endswith(".xml"):
                root = ElementTree.fromstring(package.read(name))
                features["tracked_insertions"] += len(root.findall(f".//{_W}ins"))
                features["tracked_deletions"] += len(root.findall(f".//{_W}del"))
                features["content_controls"] += len(root.findall(f".//{_W}sdt"))
                features["simple_fields"] += len(root.findall(f".//{_W}fldSimple"))
                features["field_instructions"] += len(root.findall(f".//{_W}instrText"))
                features["comment_references"] += len(
                    root.findall(f".//{_W}commentReference")
                )
            elif name.endswith(".rels"):
                root = ElementTree.fromstring(package.read(name))
                for relationship in root.findall(f"{_REL}Relationship"):
                    if relationship.get("TargetMode") == "External":
                        external_links.append(relationship.get("Target", ""))

    return {
        "path": str(path),
        "size_bytes": path.stat().st_size,
        "paragraphs": len(document.paragraphs),
        "nonempty_paragraphs": len(paragraphs),
        "paragraph_sample": [
            {"style": paragraph.style.name, "text": paragraph.text}
            for paragraph in paragraphs[:sample_paragraphs]
        ],
        "headings": headings,
        "tables": tables,
        "sections": sections,
        "inline_shapes": len(document.inline_shapes),
        "comments": len(document.comments),
        "headers": sum(name.startswith("word/header") for name in names),
        "footers": sum(name.startswith("word/footer") for name in names),
        "footnotes": "word/footnotes.xml" in names,
        "endnotes": "word/endnotes.xml" in names,
        "custom_properties": "docProps/custom.xml" in names,
        "contains_vba": "word/vbaProject.bin" in names,
        "embedded_objects": [
            name
            for name in names
            if name.startswith("word/embeddings/") and not name.endswith("/")
        ],
        "external_links": external_links,
        "features": features,
        "core_properties": {
            "title": properties.title,
            "subject": properties.subject,
            "author": properties.author,
            "last_modified_by": properties.last_modified_by,
            "keywords": properties.keywords,
            "created": properties.created,
            "modified": properties.modified,
        },
    }


def main() -> int:
    """Run the DOCX inspector.

    Returns:
        Process exit code.
    """
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="DOCX input")
    parser.add_argument("--sample-paragraphs", type=int, default=12)
    args = parser.parse_args()
    path = args.input.expanduser().resolve(strict=False)
    if not path.is_file():
        parser.error(f"Input file not found: {path}")
    if path.suffix.lower() != ".docx":
        parser.error("Input path must end with .docx")
    if args.sample_paragraphs < 0:
        parser.error("Sample limit must be non-negative.")
    try:
        result = inspect_docx(path, args.sample_paragraphs)
    except Exception as exc:
        parser.error(str(exc))
    print(json.dumps(result, ensure_ascii=False, indent=2, default=str))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
