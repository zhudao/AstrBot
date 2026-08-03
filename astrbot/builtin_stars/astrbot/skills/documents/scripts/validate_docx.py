#!/usr/bin/env python3
"""Validate DOCX package structure and related OOXML parts."""

from __future__ import annotations

import argparse
import json
import posixpath
import zipfile
from pathlib import Path, PurePosixPath
from xml.etree import ElementTree

from docx import Document

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_REL = "{http://schemas.openxmlformats.org/package/2006/relationships}"
_REQUIRED_PARTS = {"[Content_Types].xml", "_rels/.rels", "word/document.xml"}


def _relationship_base(rels_name: str) -> str:
    """Return the package base directory for a relationships part.

    Args:
        rels_name: Relationships-part path.

    Returns:
        Directory used to resolve internal targets.
    """
    path = PurePosixPath(rels_name)
    if rels_name == "_rels/.rels":
        return ""
    parent = path.parent
    if parent.name != "_rels":
        return str(parent)
    return str(parent.parent)


def validate_docx(path: Path) -> dict:
    """Validate a DOCX package.

    Args:
        path: DOCX input.

    Returns:
        Validation report.
    """
    errors: list[str] = []
    warnings: list[str] = []
    statistics = {
        "xml_parts": 0,
        "relationships": 0,
        "external_relationships": 0,
        "comments": 0,
        "tracked_insertions": 0,
        "tracked_deletions": 0,
        "content_controls": 0,
    }
    comment_ids: set[str] = set()
    comment_reference_ids: set[str] = set()

    with zipfile.ZipFile(path) as package:
        bad_member = package.testzip()
        if bad_member:
            errors.append(f"Corrupt ZIP member: {bad_member}")
        names = set(package.namelist())
        for required in sorted(_REQUIRED_PARTS - names):
            errors.append(f"Required package part is missing: {required}")

        for name in sorted(names):
            if not (name.endswith(".xml") or name.endswith(".rels")):
                continue
            statistics["xml_parts"] += 1
            try:
                root = ElementTree.fromstring(package.read(name))
            except ElementTree.ParseError as exc:
                errors.append(f"Invalid XML in {name}: {exc}")
                continue

            if name.endswith(".rels"):
                base = _relationship_base(name)
                for relationship in root.findall(f"{_REL}Relationship"):
                    statistics["relationships"] += 1
                    target = relationship.get("Target", "")
                    if relationship.get("TargetMode") == "External":
                        statistics["external_relationships"] += 1
                        continue
                    resolved = posixpath.normpath(
                        posixpath.join(base, target.lstrip("/"))
                    )
                    if resolved not in names:
                        errors.append(
                            f"Missing relationship target from {name}: {target}"
                        )
                continue

            statistics["tracked_insertions"] += len(root.findall(f".//{_W}ins"))
            statistics["tracked_deletions"] += len(root.findall(f".//{_W}del"))
            statistics["content_controls"] += len(root.findall(f".//{_W}sdt"))
            if name == "word/comments.xml":
                comment_ids.update(
                    comment.get(f"{_W}id", "")
                    for comment in root.findall(f"{_W}comment")
                )
            elif name.startswith("word/"):
                comment_reference_ids.update(
                    reference.get(f"{_W}id", "")
                    for reference in root.findall(f".//{_W}commentReference")
                )

    statistics["comments"] = len(comment_ids)
    missing_comments = comment_reference_ids - comment_ids
    if missing_comments:
        errors.append(
            "Comment references have no matching comment: "
            + ", ".join(sorted(missing_comments))
        )
    orphan_comments = comment_ids - comment_reference_ids
    if orphan_comments:
        warnings.append(
            "Comments have no visible reference anchor: "
            + ", ".join(sorted(orphan_comments))
        )

    try:
        document = Document(path)
        if not document.sections:
            errors.append("Document contains no section.")
        for index, section in enumerate(document.sections):
            if not section.page_width or not section.page_height:
                errors.append(f"Section {index} has an invalid page size.")
            for name, margin in (
                ("top", section.top_margin),
                ("right", section.right_margin),
                ("bottom", section.bottom_margin),
                ("left", section.left_margin),
            ):
                if margin is not None and margin < 0:
                    errors.append(f"Section {index} has a negative {name} margin.")
        statistics["paragraphs"] = len(document.paragraphs)
        statistics["tables"] = len(document.tables)
        statistics["inline_shapes"] = len(document.inline_shapes)
    except Exception as exc:
        errors.append(f"python-docx could not open the package: {exc}")

    return {
        "valid": not errors,
        "path": str(path),
        "statistics": statistics,
        "errors": errors,
        "warnings": warnings,
    }


def main() -> int:
    """Run DOCX validation.

    Returns:
        Zero for a valid document, otherwise one.
    """
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="DOCX input")
    args = parser.parse_args()
    path = args.input.expanduser().resolve(strict=False)
    if not path.is_file():
        parser.error(f"Input file not found: {path}")
    if path.suffix.lower() != ".docx":
        parser.error("Input path must end with .docx")
    try:
        result = validate_docx(path)
    except Exception as exc:
        result = {
            "valid": False,
            "path": str(path),
            "statistics": {},
            "errors": [str(exc)],
            "warnings": [],
        }
    print(json.dumps(result, ensure_ascii=False, indent=2, default=str))
    return 0 if result["valid"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
