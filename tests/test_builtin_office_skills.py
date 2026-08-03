from __future__ import annotations

import csv
import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook

BUILTIN_SKILLS_DIR = (
    Path(__file__).resolve().parents[1]
    / "astrbot"
    / "builtin_stars"
    / "astrbot"
    / "skills"
)
SPREADSHEET_SCRIPTS = BUILTIN_SKILLS_DIR / "spreadsheets" / "scripts"
DOCUMENT_SCRIPTS = BUILTIN_SKILLS_DIR / "documents" / "scripts"


def _run_script(
    script: Path, *arguments: Path | str
) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(script), *(str(argument) for argument in arguments)],
        check=False,
        capture_output=True,
        text=True,
    )


def test_spreadsheet_skill_converts_inspects_and_validates_csv(tmp_path: Path) -> None:
    source = tmp_path / "source.csv"
    output = tmp_path / "output.xlsx"
    source.write_text(
        "ID,Name,Amount,Date,Formula-like text\n"
        '001,测试,1200.5,2026-08-01,"=2+2"\n'
        "002,AstrBot,875,2026-08-02,plain\n",
        encoding="utf-8",
    )

    converted = _run_script(SPREADSHEET_SCRIPTS / "csv_to_xlsx.py", source, output)

    assert converted.returncode == 0, converted.stderr
    workbook = load_workbook(output, data_only=False)
    sheet = workbook["Data"]
    assert sheet.freeze_panes == "A2"
    assert sheet["A2"].value == "001"
    assert sheet["C2"].value == 1200.5
    assert sheet["D2"].value.date().isoformat() == "2026-08-01"
    assert sheet["E2"].value == "=2+2"
    assert sheet["E2"].data_type == "s"
    assert "ImportedData" in sheet.tables
    assert not sheet.sheet_view.showGridLines
    workbook.close()

    inspected = _run_script(SPREADSHEET_SCRIPTS / "inspect_workbook.py", output)
    validated = _run_script(SPREADSHEET_SCRIPTS / "validate_workbook.py", output)

    assert inspected.returncode == 0, inspected.stderr
    inspection = json.loads(inspected.stdout)
    assert inspection["kind"] == "openxml-workbook"
    assert inspection["sheets"][0]["name"] == "Data"
    assert inspection["tables"] == 1
    assert validated.returncode == 0, validated.stdout
    assert json.loads(validated.stdout)["valid"] is True


def test_spreadsheet_skill_rejects_broken_formula_reference(tmp_path: Path) -> None:
    path = tmp_path / "broken.xlsx"
    workbook = Workbook()
    workbook.active["A1"] = "=#REF!+1"
    workbook.save(path)

    validated = _run_script(SPREADSHEET_SCRIPTS / "validate_workbook.py", path)

    assert validated.returncode == 1
    report = json.loads(validated.stdout)
    assert report["valid"] is False
    assert any("Broken formula reference" in error for error in report["errors"])


def test_document_skill_inspects_validates_scrubs_and_extracts_table(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.docx"
    sanitized = tmp_path / "sanitized.docx"
    table_csv = tmp_path / "table.csv"
    document = Document()
    document.core_properties.author = "Private Author"
    document.core_properties.last_modified_by = "Private Editor"
    document.add_heading("AstrBot 文档测试", level=1)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("这是带有批注的正文。")
    document.add_comment(run, text="请复核这句话", author="Reviewer", initials="RV")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "状态"
    table.cell(1, 0).text = "DOCX Skill"
    table.cell(1, 1).text = "完成"
    document.save(source)

    inspected = _run_script(DOCUMENT_SCRIPTS / "inspect_docx.py", source)
    validated = _run_script(DOCUMENT_SCRIPTS / "validate_docx.py", source)
    scrubbed = _run_script(DOCUMENT_SCRIPTS / "privacy_scrub.py", source, sanitized)
    extracted = _run_script(
        DOCUMENT_SCRIPTS / "docx_table_to_csv.py", source, table_csv
    )

    assert inspected.returncode == 0, inspected.stderr
    inspection = json.loads(inspected.stdout)
    assert inspection["comments"] == 1
    assert inspection["tables"][0]["rows"] == 2
    assert inspection["headings"][0]["text"] == "AstrBot 文档测试"
    assert validated.returncode == 0, validated.stdout
    assert json.loads(validated.stdout)["valid"] is True
    assert scrubbed.returncode == 0, scrubbed.stderr
    clean_document = Document(sanitized)
    assert not clean_document.core_properties.author
    assert not clean_document.core_properties.last_modified_by
    assert next(iter(clean_document.comments)).author == "Author"
    assert extracted.returncode == 0, extracted.stderr
    with table_csv.open(encoding="utf-8-sig", newline="") as csv_file:
        assert list(csv.reader(csv_file)) == [
            ["项目", "状态"],
            ["DOCX Skill", "完成"],
        ]


def test_office_skill_docs_cover_advanced_workflows() -> None:
    spreadsheets = (BUILTIN_SKILLS_DIR / "spreadsheets" / "SKILL.md").read_text(
        encoding="utf-8"
    )
    documents = (BUILTIN_SKILLS_DIR / "documents" / "SKILL.md").read_text(
        encoding="utf-8"
    )

    for text in ("keep_vba=True", "conditional formatting", "data validation"):
        assert text in spreadsheets
    for text in ("tracked revisions", "content controls", "digital signature"):
        assert text in documents
    assert "Codex" not in spreadsheets
    assert "Codex" not in documents
